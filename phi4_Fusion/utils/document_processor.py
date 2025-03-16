"""
document_processor.py - Document processing utilities

This module provides utilities for processing legal documents in various formats,
extracting text, and preparing them for analysis.

"""

import logging
import os
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union, Any

import numpy as np


class DocumentProcessor:
    """
    Document processor for handling various file formats.
    
    Supported formats:
    - Plain text (.txt)
    - PDF (.pdf)
    - Word (.docx, .doc)
    - HTML (.html, .htm)
    """
    
    def __init__(self, 
                extract_tables: bool = True,
                extract_images: bool = False,
                ocr_enabled: bool = False):
        """
        Initialize the document processor.
        
        Args:
            extract_tables: Whether to extract tables from documents
            extract_images: Whether to extract images from documents
            ocr_enabled: Whether to use OCR for scanned documents
        """
        self.extract_tables = extract_tables
        self.extract_images = extract_images
        self.ocr_enabled = ocr_enabled
        
        # Load dependencies lazily when needed
        self._pdf_extractor = None
        self._docx_extractor = None
        self._html_extractor = None
        self._ocr_engine = None
        
        logging.info("Document processor initialized")
    
    def _get_pdf_extractor(self):
        """Lazy-load PDF extraction library."""
        if self._pdf_extractor is None:
            try:
                import pypdf
                self._pdf_extractor = pypdf
            except ImportError:
                logging.error("pypdf not installed. Please install with: pip install pypdf")
                raise
        return self._pdf_extractor
    
    def _get_docx_extractor(self):
        """Lazy-load Word document extraction library."""
        if self._docx_extractor is None:
            try:
                import docx
                self._docx_extractor = docx
            except ImportError:
                logging.error("python-docx not installed. Please install with: pip install python-docx")
                raise
        return self._docx_extractor
    
    def _get_html_extractor(self):
        """Lazy-load HTML extraction library."""
        if self._html_extractor is None:
            try:
                from bs4 import BeautifulSoup
                self._html_extractor = BeautifulSoup
            except ImportError:
                logging.error("BeautifulSoup not installed. Please install with: pip install beautifulsoup4")
                raise
        return self._html_extractor
    
    def _get_ocr_engine(self):
        """Lazy-load OCR engine."""
        if self._ocr_engine is None and self.ocr_enabled:
            try:
                import pytesseract
                self._ocr_engine = pytesseract
            except ImportError:
                logging.error("pytesseract not installed. Please install with: pip install pytesseract")
                raise
        return self._ocr_engine
    
    def process_document(self, file_path: str) -> str:
        """
        Process a document file and extract its text.
        
        Args:
            file_path: Path to document file
            
        Returns:
            Extracted text content
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Get file extension
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # Process based on file type
        if file_ext == ".txt":
            return self._process_text_file(file_path)
        elif file_ext == ".pdf":
            return self._process_pdf_file(file_path)
        elif file_ext in [".docx", ".doc"]:
            return self._process_word_file(file_path)
        elif file_ext in [".html", ".htm"]:
            return self._process_html_file(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _process_text_file(self, file_path: str) -> str:
        """Process plain text file."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
            return text
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, "r", encoding="latin-1") as f:
                text = f.read()
            return text
    
    def _process_pdf_file(self, file_path: str) -> str:
        """Process PDF file."""
        extractor = self._get_pdf_extractor()
        
        try:
            pdf = extractor.PdfReader(file_path)
            
            # Extract text from each page
            text_parts = []
            for page in pdf.pages:
                text_parts.append(page.extract_text())
            
            # Check if we need OCR
            full_text = "\n".join(text_parts)
            if len(full_text.strip()) < 100 and self.ocr_enabled:
                logging.info(f"PDF appears to be scanned, using OCR: {file_path}")
                return self._process_with_ocr(file_path)
            
            return full_text
        except Exception as e:
            logging.error(f"Error processing PDF file {file_path}: {e}")
            if self.ocr_enabled:
                logging.info("Attempting to use OCR as fallback")
                return self._process_with_ocr(file_path)
            raise
    
    def _process_word_file(self, file_path: str) -> str:
        """Process Word document."""
        extractor = self._get_docx_extractor()
        
        try:
            doc = extractor.Document(file_path)
            
            # Extract text from paragraphs
            text_parts = []
            for para in doc.paragraphs:
                text_parts.append(para.text)
            
            # Extract tables if enabled
            if self.extract_tables:
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text for cell in row.cells])
                        text_parts.append(row_text)
            
            return "\n".join(text_parts)
        except Exception as e:
            logging.error(f"Error processing Word file {file_path}: {e}")
            raise
    
    def _process_html_file(self, file_path: str) -> str:
        """Process HTML file."""
        BeautifulSoup = self._get_html_extractor()
        
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                html_content = f.read()
            
            # Parse HTML
            soup = BeautifulSoup(html_content, "html.parser")
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.extract()
            
            # Get text
            text = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = "\n".join(chunk for chunk in chunks if chunk)
            
            return text
        except Exception as e:
            logging.error(f"Error processing HTML file {file_path}: {e}")
            raise
    
    def _process_with_ocr(self, file_path: str) -> str:
        """Process document using OCR."""
        if not self.ocr_enabled:
            raise RuntimeError("OCR is not enabled")
        
        ocr = self._get_ocr_engine()
        
        try:
            from PIL import Image
            import pytesseract
            
            # For PDF, we need to convert to images first
            if file_path.lower().endswith(".pdf"):
                import fitz  # PyMuPDF
                
                doc = fitz.open(file_path)
                text_parts = []
                
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    pix = page.get_pixmap()
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    text_parts.append(pytesseract.image_to_string(img))
                
                return "\n".join(text_parts)
            else:
                # For image files, process directly
                text = pytesseract.image_to_string(Image.open(file_path))
                return text
                
        except Exception as e:
            logging.error(f"Error processing document with OCR {file_path}: {e}")
            raise
    
    def extract_legal_citations(self, text: str) -> List[Dict[str, str]]:
        """
        Extract legal citations from text.
        
        Args:
            text: Document text
            
        Returns:
            List of dictionaries with citation information
        """
        # Define patterns for common legal citations
        patterns = [
            # Case citation pattern
            r"([A-Z][a-z]+\s+v\.\s+[A-Z][a-z]+)[\s,](\d+)\s+([A-Za-z\.]+)\s+(\d+)[\s,]\s*\((\d{4})\)",
            # Statute pattern
            r"(\d+)\s+([A-Z]\.[A-Z]\.)\s+§§?\s+(\d+[a-z]?(?:\([a-z0-9]+\))?)",
            # Regulation pattern
            r"(\d+)\s+C\.F\.R\.\s+§§?\s+(\d+\.\d+)"
        ]
        
        citations = []
        
        for pattern in patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                citation = {
                    "text": match.group(0),
                    "span": (match.start(), match.end())
                }
                citations.append(citation)
        
        return citations
    
    def segment_document(self, text: str, min_segment_length: int = 100) -> List[Dict[str, Any]]:
        """
        Segment a document into logical sections.
        
        Args:
            text: Document text
            min_segment_length: Minimum length for a segment
            
        Returns:
            List of dictionaries with segment information
        """
        # Simple segmentation by headers and paragraphs
        segments = []
        
        # Pattern for potential headers
        header_pattern = r"(?:^|\n)([A-Z][A-Z\s]+:?|\d+\.\s+[A-Z][a-zA-Z\s]+:?)(?:\n|\s)"
        
        # Find potential headers
        header_matches = list(re.finditer(header_pattern, text))
        
        # If no headers found, just split by paragraphs
        if not header_matches:
            paragraphs = [p for p in text.split("\n\n") if len(p.strip()) >= min_segment_length]
            
            for i, para in enumerate(paragraphs):
                segments.append({
                    "id": f"p{i}",
                    "text": para,
                    "type": "paragraph"
                })
                
            return segments
        
        # Process segments based on headers
        for i in range(len(header_matches)):
            start = header_matches[i].start()
            
            # For all except the last header
            if i < len(header_matches) - 1:
                end = header_matches[i + 1].start()
            else:
                end = len(text)
            
            segment_text = text[start:end]
            header_text = header_matches[i].group(1).strip()
            
            if len(segment_text) >= min_segment_length:
                segments.append({
                    "id": f"h{i}",
                    "text": segment_text,
                    "type": "section",
                    "header": header_text
                })
        
        return segments


"""
SUMMARY:
- Provides utilities for processing legal documents in various formats
- Supports text, PDF, Word, and HTML documents
- Includes OCR capabilities for scanned documents
- Implements legal citation extraction
- Performs document segmentation for better retrieval

TODO:
- Add support for image extraction and analysis
- Enhance table extraction with structure preservation
- Implement more citation formats and jurisdictions
- Add document classification capabilities
- Support for contract-specific parsing
- Implement legal-specific document segmentation
- Add metadata extraction from document properties
"""
